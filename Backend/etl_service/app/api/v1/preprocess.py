"""Preprocessing utilities for CV text extraction and cleaning."""
import os
import time
import unicodedata
import re
from io import BytesIO
import fitz
import pdfplumber
from docx import Document
from geopy.geocoders import Nominatim
from geopy.exc import GeocoderTimedOut
from genderize import Genderize

# Extract text from various document formats
# Currently supports PDF, DOCX, and TXT files
# ---------- PDF extractor (adjusted for file-like objects) ----------
def extract_text_from_pdf(file_obj):
    """Extract text and links from a PDF file-like object."""

    final_text = ""
    all_links = set()

    # Ensure file pointer is at start
    file_obj.seek(0)

    # Read file into bytes
    file_bytes = file_obj.read()
    pdf_stream = BytesIO(file_bytes)

    # PyMuPDF for links
    doc = fitz.open(stream=pdf_stream, filetype="pdf")
    for page in doc:
        for link in page.get_links():
            uri = link.get("uri")
            if uri and uri.startswith("http"):
                all_links.add(uri.strip())
    doc.close()

    # Reset stream for pdfplumber
    pdf_stream.seek(0)
    with pdfplumber.open(pdf_stream) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""

            url_pattern = re.compile(
                r'(https?://[^\s]+|www\.[^\s]+)',
                re.IGNORECASE
            )
            found_urls = url_pattern.findall(text)
            for u in found_urls:
                all_links.add(u.strip(").,;:!?"))

            text = re.sub(r"\(cid:\d+\)", "", text)
            text = re.sub(r"\s+", " ", text)

            if i == 1 and all_links:
                top_links = [link for link in all_links if any(k in link.lower() for k in [
                    "linkedin","github","portfolio","vercel","netlify",
                    "streamlit","huggingface","render","demo","live","project"
                ])]
                if top_links:
                    text = text.strip() + "\n\nLinks: " + ", ".join(sorted(top_links))

            final_text += f"\n\n--- Page {i} ---\n{text.strip()}"

    return final_text.strip()


# ---------- DOCX extractor ----------
def extract_text_from_docx(file_obj):
    """Extract text from a DOCX file-like object."""

    if isinstance(file_obj, str):
        doc = Document(file_obj)
    else:  # BytesIO
        file_obj.seek(0)
        doc = Document(file_obj)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text

# ---------- General extractor ----------
def extract_text(file_obj, filename=None):
    """Extract text from a file-like object based on its extension."""

    # Determine extension
    if filename:
        ext = os.path.splitext(filename)[1].lower()
    elif isinstance(file_obj, str):
        ext = os.path.splitext(file_obj)[1].lower()
    else:
        raise ValueError("Cannot determine file extension")

    if ext == ".pdf":
        return extract_text_from_pdf(file_obj)
    if ext == ".docx":
        return extract_text_from_docx(file_obj)
    if ext == ".txt":
        file_obj.seek(0)
        return file_obj.read().decode("utf-8", errors="ignore")

    raise ValueError(f"Unsupported file format: {ext}")

#Clean text
def clean_text(text: str) -> str:
    """Clean and normalize extracted CV text for LLM extraction."""

    # Normalize Unicode and typographic symbols
    text = unicodedata.normalize("NFKC", text)
    text = re.sub(r"\(cid:\d+\)", "", text)
    text = re.sub(r"â€“|â€”|–|—", "-", text)
    text = re.sub(r"[“”]", '"', text)
    text = re.sub(r"[‘’]", "'", text)
    text = re.sub(r"•+", "•", text)
    text = re.sub(r"---\s*Page\s*\d+\s*---", " ", text)

    # Replace known icons/emojis with labels
    replacements = {
        "📍": "Location:",
        "📧": "Email:",
        "📱": "Phone:",
        "🌐": "Website:",
        "💼": "LinkedIn:",
        "🐙": "GitHub:",
        "🏠": "Address:",
        "☎️": "Phone:",
        "✉️": "Email:",
    }
    for k, v in replacements.items():
        text = text.replace(k, v)

    # Normalize bullets, newlines, and separators
    text = re.sub(r'[\u2022\u25CF\u25A0•▪]', '-', text)  # normalize bullets
    text = re.sub(r'[-_]{3,}', ' ', text)
    text = re.sub(r'\n{2,}', '\n', text)
    text = re.sub(r'[ \t]+', ' ', text)

    # Strip emojis or pictographs (catch-all)
    text = re.sub(r'[\U00010000-\U0010ffff]', ' ', text)  # remove all emojis

    # Remove stray non-ASCII junk but keep letters
    text = re.sub(r'[^\x00-\x7F]+', ' ', text)

    # Normalize punctuation & spacing
    text = re.sub(r'\s([,.!?;:])', r'\1', text)
    text = re.sub(r'\s{2,}', ' ', text)
    text = text.strip()

    return text

# Initialize geolocator with a longer timeout
geolocator = Nominatim(user_agent="cv_app", timeout=5)

def get_lat_lon(location, retries=3, delay=1):
    """
    Get latitude, longitude, and country for a location string.
    
    Args:
        location (str): Location string
        retries (int): Number of retries on timeout
        delay (int): Delay between retries in seconds

    Returns:
        tuple: (latitude, longitude, country)
    """

    default_lat, default_lon, default_country = 20.5937, 78.9629, "India"

    if not location:
        return default_lat, default_lon, default_country

    attempt = 0
    while attempt < retries:
        try:
            loc = geolocator.geocode(location, addressdetails=True)
            if loc:
                latitude = loc.latitude
                longitude = loc.longitude
                country = loc.raw.get("address", {}).get("country", default_country)
                return latitude, longitude, country

            return default_lat, default_lon, default_country

        except GeocoderTimedOut:
            attempt += 1
            time.sleep(delay)

        except Exception:
            return default_lat, default_lon, default_country

    # If all retries fail
    return default_lat, default_lon, default_country

# Find Gender
# Initialize Genderize once (to avoid repeated API calls)
genderize = Genderize()
def get_gender(name):
    """ Get gender from name using Genderize"""

    try:
        result = genderize.get([name])[0]  # Genderize expects a list
        gender = result['gender']
        if gender is None:
            return 'unknown'
        return gender
    except Exception:
        return 'unknown'
